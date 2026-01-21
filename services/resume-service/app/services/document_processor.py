"""
Document processing service for handling multiple file formats.

Supports PDF, DOCX, TXT, and image files with OCR.
Requirements: 2.1, 2.2, 2.3
"""

import io
import logging
from typing import Dict, Any, Optional
from pathlib import Path

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing various document formats"""
    
    def __init__(self):
        """Initialize document processor"""
        self.supported_formats = {'.txt'}  # Always support TXT
        
        if PDF_AVAILABLE:
            self.supported_formats.add('.pdf')
        if DOCX_AVAILABLE:
            self.supported_formats.add('.docx')
        if OCR_AVAILABLE:
            self.supported_formats.update({'.png', '.jpg', '.jpeg', '.tiff'})
            
        logger.info(f"Document processor initialized with formats: {self.supported_formats}")
    
    def process_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process document and extract text content"""
        file_ext = Path(filename).suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        try:
            if file_ext == '.pdf':
                if not PDF_AVAILABLE:
                    raise ValueError("PDF processing not available - PyPDF2 not installed")
                return self._process_pdf(file_content)
            elif file_ext == '.docx':
                if not DOCX_AVAILABLE:
                    raise ValueError("DOCX processing not available - python-docx not installed")
                return self._process_docx(file_content)
            elif file_ext == '.txt':
                return self._process_txt(file_content)
            elif file_ext in {'.png', '.jpg', '.jpeg', '.tiff'}:
                if not OCR_AVAILABLE:
                    raise ValueError("Image processing not available - pytesseract/PIL not installed")
                return self._process_image(file_content)
            else:
                raise ValueError(f"Handler not implemented for: {file_ext}")
                
        except Exception as e:
            logger.error(f"Error processing {filename}: {e}")
            return {
                "text": "",
                "metadata": {
                    "filename": filename,
                    "format": file_ext,
                    "error": str(e),
                    "success": False
                }
            }
    
    def _process_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text from PDF file"""
        text = ""
        metadata = {"format": "pdf", "pages": 0, "success": True}
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            metadata["pages"] = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    text += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {e}")
                    metadata["warnings"] = metadata.get("warnings", [])
                    metadata["warnings"].append(f"Page {page_num}: {str(e)}")
            
            # If no text extracted, might be scanned PDF
            if not text.strip():
                logger.info("No text found in PDF, might be scanned document")
                metadata["likely_scanned"] = True
                
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            metadata["success"] = False
            metadata["error"] = str(e)
        
        return {
            "text": text.strip(),
            "metadata": metadata
        }
    
    def _process_docx(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        text = ""
        metadata = {"format": "docx", "paragraphs": 0, "success": True}
        
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            text = "\n".join(paragraphs)
            metadata["paragraphs"] = len(paragraphs)
            
            # Extract tables if any
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(" | ".join(row_text))
            
            if tables_text:
                text += "\n\nTables:\n" + "\n".join(tables_text)
                metadata["tables"] = len(doc.tables)
                
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            metadata["success"] = False
            metadata["error"] = str(e)
        
        return {
            "text": text.strip(),
            "metadata": metadata
        }
    
    def _process_txt(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text from TXT file"""
        metadata = {"format": "txt", "success": True}
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = ""
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    metadata["encoding"] = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if not text:
                raise ValueError("Could not decode text file with any supported encoding")
                
        except Exception as e:
            logger.error(f"Error processing TXT: {e}")
            metadata["success"] = False
            metadata["error"] = str(e)
            text = ""
        
        return {
            "text": text.strip(),
            "metadata": metadata
        }
    
    def _process_image(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text from image using OCR"""
        text = ""
        metadata = {"format": "image", "success": True, "ocr": True}
        
        try:
            image = Image.open(io.BytesIO(file_content))
            metadata["image_size"] = image.size
            metadata["image_mode"] = image.mode
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            
            # Get OCR confidence if available
            try:
                ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
                confidences = [int(conf) for conf in ocr_data['conf'] if int(conf) > 0]
                if confidences:
                    metadata["ocr_confidence"] = sum(confidences) / len(confidences)
                else:
                    metadata["ocr_confidence"] = 0
            except Exception as e:
                logger.warning(f"Could not get OCR confidence: {e}")
                metadata["ocr_confidence"] = None
            
        except Exception as e:
            logger.error(f"Error processing image with OCR: {e}")
            metadata["success"] = False
            metadata["error"] = str(e)
        
        return {
            "text": text.strip(),
            "metadata": metadata
        }
    
    def assess_quality(self, result: Dict[str, Any]) -> Dict[str, Any]:
        """Assess the quality of extracted text"""
        text = result.get("text", "")
        metadata = result.get("metadata", {})
        
        quality_score = 1.0
        issues = []
        
        # Check text length
        if len(text) < 50:
            quality_score -= 0.3
            issues.append("Very short text extracted")
        
        # Check for OCR quality
        if metadata.get("ocr"):
            ocr_confidence = metadata.get("ocr_confidence")
            if ocr_confidence is not None:
                if ocr_confidence < 50:
                    quality_score -= 0.4
                    issues.append("Low OCR confidence")
                elif ocr_confidence < 70:
                    quality_score -= 0.2
                    issues.append("Medium OCR confidence")
        
        # Check for processing errors
        if not metadata.get("success", True):
            quality_score = 0.0
            issues.append("Processing failed")
        
        # Check for warnings
        if metadata.get("warnings"):
            quality_score -= 0.1
            issues.append("Processing warnings")
        
        quality_score = max(0.0, quality_score)
        
        return {
            "quality_score": quality_score,
            "issues": issues,
            "needs_review": quality_score < 0.7 or len(issues) > 0
        }